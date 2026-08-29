"""Focused smoke tests for dependency APIs used at production boundaries.

These are deliberately small and offline. They complement application tests by
making important third-party API assumptions fail close to a dependency update.
"""

from __future__ import annotations

import inspect
import json
import logging
from io import BytesIO

import asyncpg
import geonamescache
import httpx
import msgspec
import orjson
import pypdf
import pytest
import structlog
from cryptography.fernet import Fernet
from docx import Document
from google import genai
from google.genai import types
from hypercorn.config import Config as HypercornConfig
from PIL import Image
from quart import Quart, jsonify
from redis.asyncio import Redis
from telegram.ext import Application, CommandHandler

from app.config import Settings, settings
from app.natal.calculator import calculate_chart
from app.natal.models import BirthInput, ResolvedBirthData, TimePrecision
from app.providers.stream_types import GenerationRequest, PromptRole, PromptTurn, TextPart
from app.providers.typed_payloads import gemini_contents
from app.utils.logging_config import configure_structlog_pipeline


async def _telegram_callback(update, context) -> None:
    del update, context


def test_telegram_application_and_polling_contract() -> None:
    application = Application.builder().token("123456:test-token").build()
    application.add_handler(CommandHandler("start", _telegram_callback))

    assert application.bot.token == "123456:test-token"
    assert application.updater is not None
    polling_parameters = inspect.signature(application.updater.start_polling).parameters
    assert {"allowed_updates", "drop_pending_updates", "poll_interval", "timeout"} <= set(polling_parameters)


def test_structlog_processor_formatter_contract() -> None:
    try:
        formatter = configure_structlog_pipeline(enable_structured=True, enable_pretty=False)
        record = logging.LogRecord(
            name="dependency-boundary",
            level=logging.INFO,
            pathname="boundary.py",
            lineno=1,
            msg="hello %s",
            args=("world",),
            exc_info=None,
        )

        assert json.loads(formatter.format(record))["event"] == "hello world"
    finally:
        structlog.reset_defaults()


@pytest.mark.asyncio
async def test_google_genai_typed_request_adapter_contract() -> None:
    request = GenerationRequest(
        models=("gemini-test",),
        turns=(PromptTurn(role=PromptRole.USER, parts=(TextPart("dependency boundary"),)),),
    )

    contents = await gemini_contents(request)
    config = types.GenerateContentConfig(temperature=0.0, max_output_tokens=8)

    assert len(contents) == 1
    assert contents[0].role == "user"
    assert contents[0].parts is not None
    assert contents[0].parts[0].text == "dependency boundary"
    assert config.max_output_tokens == 8


@pytest.mark.asyncio
async def test_google_genai_manual_function_call_survives_sdk_wrapper() -> None:
    requests: list[dict[str, object]] = []

    async def handle_request(request: httpx.Request) -> httpx.Response:
        requests.append(json.loads(request.content))
        return httpx.Response(
            200,
            json={
                "candidates": [
                    {
                        "content": {
                            "role": "model",
                            "parts": [
                                {
                                    "functionCall": {
                                        "name": "search_web",
                                        "args": {"query": "Kyiv"},
                                    }
                                }
                            ],
                        },
                        "finishReason": "STOP",
                    }
                ]
            },
        )

    client = genai.Client(
        api_key="test-key",
        http_options=types.HttpOptions(async_client_args={"transport": httpx.MockTransport(handle_request)}),
    )
    config = types.GenerateContentConfig(
        tools=[
            types.Tool(
                function_declarations=[
                    types.FunctionDeclaration(
                        name="search_web",
                        description="Search the web",
                        parameters={
                            "type": "object",
                            "properties": {"query": {"type": "string"}},
                            "required": ["query"],
                        },
                    )
                ]
            )
        ]
    )

    try:
        response = await client.aio.models.generate_content(
            model="gemini-test",
            contents="Find Kyiv",
            config=config,
        )
    finally:
        await client.aio.aclose()

    assert len(requests) == 1
    assert requests[0]["tools"] == [
        {
            "functionDeclarations": [
                {
                    "description": "Search the web",
                    "name": "search_web",
                    "parameters": {
                        "properties": {"query": {"type": "STRING"}},
                        "required": ["query"],
                        "type": "OBJECT",
                    },
                }
            ]
        }
    ]
    assert response.function_calls is not None
    assert response.function_calls[0].name == "search_web"
    assert response.function_calls[0].args == {"query": "Kyiv"}


@pytest.mark.asyncio
async def test_quart_test_client_and_hypercorn_configuration_contract() -> None:
    app = Quart(__name__)

    @app.get("/health")
    async def health():
        return jsonify(status="ok")

    response = await app.test_client().get("/health")
    config = HypercornConfig()
    config.bind = ["127.0.0.1:10000"]

    assert response.status_code == 200
    assert await response.get_json() == {"status": "ok"}
    assert config.bind == ["127.0.0.1:10000"]


def test_pydantic_application_settings_round_trip() -> None:
    cloned = Settings.model_validate(settings.model_dump())

    assert cloned.TELEGRAM_BOT_TOKEN == settings.TELEGRAM_BOT_TOKEN
    assert cloned.DATABASE_URL == settings.DATABASE_URL
    assert cloned.model_dump(mode="json") == settings.model_dump(mode="json")


def test_crypto_document_and_image_round_trips() -> None:
    fernet = Fernet(Fernet.generate_key())
    plaintext = b"dependency-boundary"
    assert fernet.decrypt(fernet.encrypt(plaintext)) == plaintext

    pdf_bytes = BytesIO()
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(pdf_bytes)
    pdf_bytes.seek(0)
    assert len(pypdf.PdfReader(pdf_bytes).pages) == 1

    docx_bytes = BytesIO()
    document = Document()
    document.add_paragraph("dependency boundary")
    document.save(docx_bytes)
    docx_bytes.seek(0)
    assert Document(docx_bytes).paragraphs[0].text == "dependency boundary"

    image_bytes = BytesIO()
    Image.new("RGB", (2, 3), color=(10, 20, 30)).save(image_bytes, format="PNG")
    image_bytes.seek(0)
    with Image.open(image_bytes) as image:
        assert image.size == (2, 3)
        assert image.getpixel((0, 0)) == (10, 20, 30)


def test_orjson_and_msgspec_typed_serialization_contract() -> None:
    class Payload(msgspec.Struct):
        name: str
        count: int

    raw = orjson.dumps({"name": "граница", "count": 2}, option=orjson.OPT_SORT_KEYS)
    decoded = msgspec.json.decode(raw, type=Payload)

    assert decoded == Payload(name="граница", count=2)
    assert orjson.loads(msgspec.json.encode(decoded)) == {"name": "граница", "count": 2}


@pytest.mark.asyncio
async def test_ephem_geonames_redis_and_asyncpg_construction_contracts() -> None:
    chart = await calculate_chart(
        ResolvedBirthData(
            birth_input=BirthInput(
                birth_date="1990-01-01",
                birth_time="12:00",
                time_precision=TimePrecision.EXACT,
                birth_place="Kyiv",
            ),
            latitude=50.4501,
            longitude=30.5234,
            timezone="Europe/Kyiv",
            local_datetime="1990-01-01T12:00:00+02:00",
            utc_datetime="1990-01-01T10:00:00Z",
            display_place="Kyiv",
        )
    )
    countries = geonamescache.GeonamesCache().get_countries()
    redis_client = Redis.from_url("redis://localhost:6379/0", decode_responses=True)
    pool_parameters = inspect.signature(asyncpg.create_pool).parameters
    connect_parameters = inspect.signature(asyncpg.connect).parameters

    assert chart.planets
    assert countries["UA"]["name"] == "Ukraine"
    assert redis_client.get_connection_kwargs()["db"] == 0
    assert {"dsn", "min_size", "max_size", "init", "connect_kwargs"} <= set(pool_parameters)
    assert {"command_timeout", "statement_cache_size", "server_settings"} <= set(connect_parameters)
    await redis_client.aclose()
