"""
Document processor — orchestration layer for document upload and processing.

The heavy lifting is delegated to submodules:
- ``app.documents.parsers``:    sync file I/O and hashing
- ``app.documents.repository``: async database CRUD and stats

This file keeps the ``DocumentProcessor`` class (orchestrator + PDF/DOCX
parsing), the singleton instance, and backward-compatible facade functions.
"""

import asyncio
import io
import logging
import tempfile
from pathlib import Path
from typing import Any

import httpx
import pypdf
from docx import Document as DocxDocument

# Submodule imports
from app.documents.parsers import (
    MAX_DOCUMENT_TEXT_LENGTH,  # noqa: F401 — re-exported
    calculate_file_hash_sync,
)
from app.documents.repository import (
    check_document_limit,
    check_duplicate_file,
    cleanup_oldest_documents,
    save_document_content,
)
from app.metrics import metrics_collector
from app.utils.network import NetworkErrorHandler

# Verify document processing libraries are available
try:
    import pypdf as _pypdf_check  # noqa: F811,F401
    from docx import Document as _docx_check  # noqa: F811,F401

    DOCUMENT_SUPPORT = True
except ImportError:
    DOCUMENT_SUPPORT = False
    logging.warning("Document processing libraries not installed. Document support disabled.")


class DocumentProcessor:
    """Orchestrates document upload, validation, parsing, and persistence."""

    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".doc"]
        self.max_file_size = 10 * 1024 * 1024  # 10MB for free tier
        self.max_pages = 50

    async def process_document(self, file_data, filename: str, user_id: int, is_path: bool = False) -> dict[str, Any]:
        """Process a document and return extracted text."""
        if not DOCUMENT_SUPPORT:
            return {"error": "Document processing is not available"}

        try:
            # Check file size
            if is_path:
                import os

                file_size = os.path.getsize(file_data)
            else:
                file_size = len(file_data)

            if file_size > self.max_file_size:
                return {"error": f"File too large. Maximum size is {self.max_file_size // (1024 * 1024)}MB"}

            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.supported_formats:
                return {"error": f"Unsupported file format: {file_ext}"}

            # Check document limit
            if not await check_document_limit(user_id):
                await cleanup_oldest_documents(user_id, 4)
                logging.info("Document limit exceeded for user %s, removed oldest document", user_id)

            # Hash + duplicate check
            loop = asyncio.get_running_loop()
            file_hash = await loop.run_in_executor(None, calculate_file_hash_sync, file_data)
            duplicate = await check_duplicate_file(user_id, file_hash, filename)

            if duplicate:
                created_date = duplicate["created_at"]
                if hasattr(created_date, "strftime"):
                    date_str = created_date.strftime("%Y-%m-%d")
                else:
                    date_str = str(created_date)[:10]

                return {
                    "error": "duplicate",
                    "message": f"Файл '{filename}' уже был загружен ранее как '{duplicate['filename']}' ({date_str})",
                    "duplicate_info": duplicate,
                }

            # Dispatch to format-specific parser
            if file_ext == ".pdf":
                return await self._process_pdf_unified(file_data, filename, user_id, file_hash, is_path=is_path)
            elif file_ext in [".docx", ".doc"]:
                return await self._process_word_unified(file_data, filename, user_id, file_hash, is_path=is_path)
            else:
                return {"error": f"Unsupported file format: {file_ext}"}

        except (ValueError, UnicodeDecodeError, OSError) as e:
            logging.error("Error processing document %s: %s", filename, e, exc_info=True)
            await metrics_collector.record_error("document_processing", str(e))
            return {"error": f"Error processing document: {str(e)}"}

    async def process_document_force(
        self, file_data, filename: str, user_id: int, is_path: bool = False
    ) -> dict[str, Any]:
        """Process a document, ignoring duplicates."""
        if not DOCUMENT_SUPPORT:
            return {"error": "Document processing is not available"}

        try:
            if is_path:
                import os

                file_size = os.path.getsize(file_data)
            else:
                file_size = len(file_data)

            if file_size > self.max_file_size:
                return {"error": f"File too large. Maximum size is {self.max_file_size // (1024 * 1024)}MB"}

            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.supported_formats:
                return {"error": f"Unsupported file format: {file_ext}"}

            loop = asyncio.get_running_loop()
            file_hash = await loop.run_in_executor(None, calculate_file_hash_sync, file_data)

            if file_ext == ".pdf":
                return await self._process_pdf_unified(file_data, filename, user_id, file_hash, is_path=is_path)
            elif file_ext in [".docx", ".doc"]:
                return await self._process_word_unified(file_data, filename, user_id, file_hash, is_path=is_path)
            else:
                return {"error": f"Unsupported file format: {file_ext}"}

        except (ValueError, UnicodeDecodeError, OSError) as e:
            logging.error("Error processing document %s: %s", filename, e, exc_info=True)
            await metrics_collector.record_error("document_processing", str(e))
            return {"error": f"Error processing document: {str(e)}"}

    # ── Synchronous parsing helpers (run via run_in_executor) ─────────────

    @staticmethod
    def _process_pdf_sync(input_data: str | io.BytesIO, max_pages: int) -> dict[str, Any]:
        """Synchronous PDF text extraction — runs in a thread executor."""
        pdf_file = None
        should_close = False

        try:
            if isinstance(input_data, str):
                pdf_file = open(input_data, "rb")  # noqa: SIM115
                should_close = True
                stream = pdf_file
            else:
                stream = input_data  # type: ignore[assignment]  # BytesIO is compatible with PdfReader

            pdf_reader = pypdf.PdfReader(stream)

            if len(pdf_reader.pages) > max_pages:
                return {"error": f"PDF too large. Maximum {max_pages} pages allowed"}

            text_content: list[str] = []
            current_length = 0

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        chunk = f"--- Page {page_num + 1} ---\n{text}"
                        text_content.append(chunk)
                        current_length += len(chunk)
                except Exception as page_error:
                    logging.warning(
                        "Error extracting text from page %d: %s",
                        page_num + 1,
                        page_error,
                    )
                    chunk = f"--- Page {page_num + 1} ---\n[Error extracting text from this page]"
                    text_content.append(chunk)
                    current_length += len(chunk)

                if current_length > MAX_DOCUMENT_TEXT_LENGTH:
                    text_content.append(f"\n--- Document truncated at page {page_num + 1} ---")
                    break

            full_text = "\n\n".join(text_content)
            return {
                "success": True,
                "pages": len(pdf_reader.pages),
                "content": full_text,
            }
        except Exception as e:
            return {"error": str(e)}
        finally:
            if should_close and pdf_file:
                pdf_file.close()

    @staticmethod
    def _process_word_sync(input_data: str | io.BytesIO) -> dict[str, Any]:
        """Synchronous Word text extraction — runs in a thread executor."""
        try:
            doc = DocxDocument(input_data)

            text_content: list[str] = []
            paragraph_count = 0

            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
                    paragraph_count += 1

            table_count = 0
            for table in doc.tables:
                table_count += 1
                text_content.append(f"\n--- Table {table_count} ---")
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))

            full_text = "\n\n".join(text_content)
            return {
                "success": True,
                "pages": 1,
                "paragraphs": paragraph_count,
                "tables": table_count,
                "text_length": len(full_text),
                "content": full_text,
            }
        except Exception as e:
            return {"error": str(e)}

    # ── PDF parsing ──────────────────────────────────────────────────────────

    async def _process_pdf_unified(
        self, file_data, filename: str, user_id: int, file_hash: str, is_path: bool = False
    ) -> dict[str, Any]:
        """Process a PDF document (bytes or path)."""
        try:
            if not is_path and not file_data.startswith(b"%PDF"):
                logging.warning("Invalid PDF format for %s", filename)
                return {"error": "Invalid PDF file format"}

            logging.info("Processing PDF %s with pypdf", filename)

            if is_path:
                sync_input = file_data
            else:
                stream = tempfile.SpooledTemporaryFile(max_size=2 * 1024 * 1024, mode="w+b")
                stream.write(file_data)
                stream.seek(0)
                sync_input = stream

            loop = asyncio.get_running_loop()
            result = await loop.run_in_executor(None, self._process_pdf_sync, sync_input, self.max_pages)

            if "error" in result:
                return result

            full_text = result["content"]
            pages_count = result["pages"]

            await save_document_content(user_id, filename, full_text, pages_count, file_hash)

            return {
                "success": True,
                "filename": filename,
                "pages": pages_count,
                "text_length": len(full_text),
                "content": full_text,
                "method": "pypdf",
            }

        except (ValueError, OSError, pypdf.errors.PdfReadError) as e:
            logging.error("Error processing PDF %s: %s", filename, e, exc_info=True)
            await metrics_collector.record_error("pdf_processing", str(e))
            return {"error": f"Error processing PDF: {str(e)}"}

    # ── Word parsing ─────────────────────────────────────────────────────────

    async def _process_word_unified(
        self, file_data, filename: str, user_id: int, file_hash: str, is_path: bool = False
    ) -> dict[str, Any]:
        """Process a Word document (bytes or path)."""
        if not is_path and not file_data.startswith(b"\x50\x4b\x03\x04"):
            logging.warning("Invalid DOCX format for %s: Missing ZIP header", filename)
            return {"error": "Invalid Word document format. File must be a valid .docx file."}

        try:
            if is_path:
                sync_input = file_data
            else:
                stream = tempfile.SpooledTemporaryFile(max_size=2 * 1024 * 1024, mode="w+b")
                stream.write(file_data)
                stream.seek(0)
                sync_input = stream

            loop = asyncio.get_running_loop()
            result = await loop.run_in_executor(None, self._process_word_sync, sync_input)

            if "error" in result:
                logging.error("Error processing Word document %s: %s", filename, result["error"])
                return {"error": f"Error processing Word document: {result['error']}"}

            full_text = result["content"]

            await save_document_content(user_id, filename, full_text, 1, file_hash)

            result["filename"] = filename
            return result

        except (ValueError, UnicodeDecodeError, OSError) as e:
            logging.error("Error processing Word document %s: %s", filename, e, exc_info=True)
            await metrics_collector.record_error("word_processing", str(e))
            return {"error": f"Error processing Word document: {str(e)}"}

    # ── Delegated repository methods (backward compat for external callers) ──

    async def _cleanup_oldest_documents(self, user_id, keep_count=4):
        return await cleanup_oldest_documents(user_id, keep_count)

    async def get_document_by_id(self, document_id, user_id):
        from app.documents.repository import get_document_by_id

        return await get_document_by_id(document_id, user_id)

    async def get_user_documents(self, user_id):
        from app.documents.repository import get_user_documents

        return await get_user_documents(user_id)

    async def get_document_content(self, document_id, user_id):
        from app.documents.repository import get_document_content

        return await get_document_content(document_id, user_id)

    async def delete_document(self, document_id, user_id):
        from app.documents.repository import delete_document

        return await delete_document(document_id, user_id)

    async def delete_all_user_documents(self, user_id):
        from app.documents.repository import delete_all_user_documents

        return await delete_all_user_documents(user_id)

    async def cleanup_old_documents(self, days_old=3):
        from app.documents.repository import cleanup_old_documents

        return await cleanup_old_documents(days_old)

    async def get_document_stats(self):
        from app.documents.repository import get_document_stats

        return await get_document_stats()

    async def get_user_document_stats(self, user_id):
        from app.documents.repository import get_user_document_stats

        return await get_user_document_stats(user_id)


# Singleton
document_processor = DocumentProcessor()


# ============================================================================
# FACADE FUNCTIONS (backward compatibility)
# ============================================================================


async def process_uploaded_document(file_data, filename: str, user_id: int, is_path: bool = False) -> dict[str, Any]:
    """Process an uploaded document."""
    return await document_processor.process_document(file_data, filename, user_id, is_path)


async def process_uploaded_document_force(
    file_data, filename: str, user_id: int, is_path: bool = False
) -> dict[str, Any]:
    """Process an uploaded document, ignoring duplicates."""
    return await document_processor.process_document_force(file_data, filename, user_id, is_path)


async def get_user_documents(user_id: int) -> list[dict[str, Any]]:
    """Get user's documents."""
    return await document_processor.get_user_documents(user_id)


async def get_document_content(document_id: int, user_id: int) -> str | None:
    """Get document content."""
    return await document_processor.get_document_content(document_id, user_id)


async def delete_user_document(document_id: int, user_id: int) -> bool:
    """Delete a user document."""
    return await document_processor.delete_document(document_id, user_id)


async def delete_all_user_documents(user_id: int) -> int:
    """Delete all user documents."""
    return await document_processor.delete_all_user_documents(user_id)


async def _upload_file_to_x0_at(file_data: bytes, filename: str) -> str | None:
    """Internal function for uploading file to x0.at with retry logic."""
    timeout_config = httpx.Timeout(
        connect=10.0,
        read=60.0,
        write=60.0,
        pool=30.0,
    )

    async with httpx.AsyncClient(timeout=timeout_config) as client:
        files = {"file": (filename, file_data)}
        response = await client.post("https://x0.at/", files=files)

        if response.status_code == 200:
            url = response.text.strip()
            if url.startswith("http"):
                logging.info("File %s uploaded to x0.at: %s", filename, url)
                return url
            else:
                logging.error("Invalid response from x0.at: %s", response.text)
                return None
        else:
            logging.error("Failed to upload to x0.at: %s - %s", response.status_code, response.text)
            return None


async def upload_to_x0_at(file_data: bytes, filename: str) -> str | None:
    """Upload file to x0.at with automatic retries."""
    try:
        return await NetworkErrorHandler.retry_with_backoff(
            _upload_file_to_x0_at,
            max_retries=3,
            base_delay=2.0,
            file_data=file_data,
            filename=filename,
        )
    except (httpx.HTTPError, OSError) as e:
        logging.error("Error uploading to x0.at after retries: %s", e, exc_info=True)
        return None


async def get_document_by_id(document_id: int, user_id: int) -> dict[str, Any] | None:
    """Get document by ID."""
    return await document_processor.get_document_by_id(document_id, user_id)
